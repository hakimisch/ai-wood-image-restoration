#!/usr/bin/env python
"""Phase 3: Fix remaining math variables with Unicode-aware matching."""

import re
from docx import Document

SRC = r'C:\Hakimi\Internship\GUI Image Restoration\research\output\paper_fixed.docx'
DST = r'C:\Hakimi\Internship\GUI Image Restoration\research\output\paper_fixed.docx'


def main():
    doc = Document(SRC)
    fixes = 0
    
    # Read ALL text to find paragraphs by content
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        if not t.strip():
            continue
        
        original = t
        text = t
        
        # === Fix 1: h_PSF — match using \u2019 (smart apostrophe) ===
        # "where is the microscope\u2019s true point spread function, a complex..."
        # -> "where h_PSF is the microscope\u2019s true point spread function, a complex..."
        if re.search(r'where is the microscope\u2019s true point spread function', text):
            text = text.replace(
                'where is the microscope\u2019s true point spread function',
                'where h_PSF is the microscope\u2019s true point spread function'
            )
            print(f'  Fixed h_PSF in paragraph [{i}]')
        
        # === Fix 2: eta_sensor — before "is read noise" ===
        # The text has "interface, is read noise, eps_banding"
        # -> "interface, eta_sensor is read noise, eps_banding"
        m = re.search(r'interface,\s+is read noise,', text)
        if m:
            old = m.group(0)
            new = old.replace('is read noise,', 'eta_sensor is read noise,')
            # But only if not already fixed
            if 'eta_sensor' not in text:
                text = text.replace(old, new)
                print(f'  Fixed eta_sensor in paragraph [{i}]')
        
        # === Fix 3: eps_vignette — "and is the radial falloff" ===
        # "modulation, and is the radial falloff attributable..."
        # -> "modulation, and eps_vignette is the radial falloff attributable..."
        m = re.search(r',\s+and is the radial falloff', text)
        if m and 'eps_vignette' not in text and 'eps_vign' not in text:
            old = m.group(0)
            new = old.replace('and is the', 'and eps_vignette is the')
            text = text.replace(old, new)
            print(f'  Fixed eps_vignette in paragraph [{i}]')
        
        # === Fix 4: Section 2.2 receptive field ===
        # "of a -layer CNN with kernels and stride grows linearly as."
        # The text has "of a -layer CNN with kernels and stride grows linearly as."
        # Note: the hyphen before "layer" is a regular hyphen
        if '-layer CNN with kernels' in text and 'D-layer' not in text:
            text = re.sub(
                r'of a\s*-layer CNN with kernels and stride grows linearly as',
                'of a D-layer CNN with K x K kernels and stride S=1 grows linearly as R = 1 + D(K-1)',
                text
            )
            print(f'  Fixed receptive field in paragraph [{i}]')
        
        # === Fix 5: Section 1.4 VoL ===
        # "where is the discrete Laplacian approximation using a kernel"
        m = re.search(r'where is the discrete Laplacian approximation using a kernel', text)
        if m and 'L is the' not in text:
            text = text.replace(
                'where is the discrete Laplacian approximation using a kernel',
                'where L (Laplacian operator) is the discrete Laplacian approximation using a 3x3 kernel'
            )
            print(f'  Fixed VoL in paragraph [{i}]')
        
        # === Fix 6: Section 3.4.1 loss function ===
        # "where is the mean squared error" -> "where L_MSE is the mean squared error"
        m = re.search(r'where is the mean squared error', text)
        if m and 'L_MSE' not in text:
            text = text.replace('where is the mean squared error', 'where L_MSE is the mean squared error')
            print(f'  Fixed L_MSE in paragraph [{i}]')
        
        # "and is the mean absolute error" -> "and L_L1 is the mean absolute error"
        m = re.search(r'and is the mean absolute error', text)
        if m and 'L_L1' not in text:
            text = text.replace('and is the mean absolute error', 'and L_L1 is the mean absolute error')
            print(f'  Fixed L_L1 in paragraph [{i}]')
        
        # === Fix 7: Section 4.3.1 norm ===
        m = re.search(r'minimise the expected norm', text)
        if m and 'L2' not in text:
            text = text.replace('minimise the expected norm', 'minimise the expected L2 norm')
            print(f'  Fixed L2 norm in paragraph [{i}]')
        
        # === Fix 8: Section 4.3.4 norm ===
        m = re.search(r'penalises the norm of the image gradient', text)
        if m and 'L1' not in text:
            text = text.replace('penalises the norm of the image gradient', 'penalises the L1 norm of the image gradient')
            print(f'  Fixed L1 norm in paragraph [{i}]')
        
        # === Fix 9: Section 4.3.4 weights ===
        m = re.search(r'weights ranging from to', text)
        if m:
            text = text.replace('weights ranging from to', 'weights ranging from 10^-4 to 10^-2')
            print(f'  Fixed TV weights in paragraph [{i}]')
        
        # === Fix 10: inverts/not ===
        m = re.search(r'invert\s*, not', text)
        if m and 'Phi_simple' not in text:
            text = re.sub(r'invert\s*, not\s*', 'invert Phi_simple, not Phi_phys', text)
            print(f'  Fixed invert in paragraph [{i}]')
        
        m = re.search(r'inverts\s*, not', text)
        if m and 'Phi_simple' not in text:
            text = re.sub(r'inverts\s*, not\s*', 'inverts Phi_simple, not Phi_phys', text)
            print(f'  Fixed inverts in paragraph [{i}]')
        
        if text != original:
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = text
            else:
                p.add_run(text)
            fixes += 1
    
    doc.save(DST)
    print(f'\nTotal fixes: {fixes}')


if __name__ == '__main__':
    main()
