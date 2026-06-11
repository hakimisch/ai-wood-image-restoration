#!/usr/bin/env python
"""
Refine the paper DOCX:
  1. Remove all em dashes (—) with context-appropriate replacements
  2. Polish writing (tighter prose, consistent terminology)
  3. Polish formatting (blank line cleanup, heading refinement)
"""

import re
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = r'C:\Hakimi\Internship\GUI Image Restoration\research\output\paper.docx'
DST = r'C:\Hakimi\Internship\GUI Image Restoration\research\output\paper.docx'


def replace_em_dashes(text, is_heading=False):
    """
    Replace em dashes (—) with context-appropriate alternatives.
    
    For headings: em dashes in subtitles become colons.
    For body text:
      - Paired em dashes (two in same region) for parentheticals -> commas
      - Single em dash before list/explanation -> colon
      - Single em dash for emphasis -> semicolon or comma
    """
    if '—' not in text:
        return text
    
    if is_heading:
        # Headings: replace em dash with colon (subtitle separator)
        text = text.replace('—', ': ')
        text = re.sub(r':\s+:', ':', text)  # remove double colons
        return text
    
    # For body text, process by sentence
    sentences = re.split(r'(?<=[.!?])\s+', text)
    result = []
    
    for sent in sentences:
        if '—' not in sent:
            result.append(sent)
            continue
        
        # Special case: "—" at end of line or followed by URL/code
        # Count em dashes
        parts = sent.split('—')
        
        # Pattern A: Paired dashes (parenthetical) — must be at least 2
        if len(parts) >= 3:
            # Process as parenthetical: odd positions are opening, even are closing
            # Replace with commas for non-restrictive clauses
            new_parts = []
            for j, part in enumerate(parts):
                new_parts.append(part)
                if j < len(parts) - 1:
                    new_parts.append(', ')
            sent = ''.join(new_parts).rstrip(', ')
            sent = re.sub(r',\s*,', ',', sent)
        
        # Pattern B/C: Single em dash
        elif len(parts) == 2:
            left = parts[0].strip()
            right = parts[1].strip()
            
            # Check if this is introducing a list or enumeration
            list_signals = [
                left.rstrip().endswith((': ', ':', ';'))
            ]
            
            # Words that typically introduce lists
            left_last_word = left.split()[-1] if left.split() else ''
            list_intro_words = ['including', 'such as', 'namely', 'following', 
                                'spanning', 'comprising', 'encompassing']
            
            is_list_intro = left_last_word.rstrip(',').lower() in list_intro_words
            
            # Check if right side looks like a list
            is_list = ',' in right and not re.match(r'^(that|which|who)', right)
            
            if is_list_intro or (is_list and not right.startswith(('that', 'which', 'who'))):
                sent = f'{left}: {right}'
            elif right[0].isupper() and len(right) > 5:
                # Appended independent clause -> semicolon
                sent = f'{left}; {right}'
            else:
                # General appositive/continuation -> comma
                sent = f'{left}, {right}'
        
        # Clean up double spaces
        sent = re.sub(r'  +', ' ', sent)
        result.append(sent)
    
    return ' '.join(result)


def polish_writing(text):
    """Apply writing refinements for tighter prose."""
    if not text.strip():
        return text
    
    # --- Verbose phrase trimming ---
    text = re.sub(r'\bIt is important to note that\b', 'Importantly,', text)
    text = re.sub(r'\bIt should be noted that\b', 'Notably,', text)
    text = re.sub(r'\bin order to\b', 'to', text)
    text = re.sub(r'\bas a result of\b', 'from', text)
    text = re.sub(r'\bin the context of\b', 'in', text)
    text = re.sub(r'\bwith respect to\b', 'regarding', text)
    text = re.sub(r'\bon the basis of\b', 'from', text)
    text = re.sub(r'\bdue to the fact that\b', 'because', text)
    text = re.sub(r'\bin spite of the fact that\b', 'although', text)
    
    # --- "It is / there are" constructions ---
    text = re.sub(r'\bThere are\s+', 'We identify ', text)
    text = re.sub(r'\bThere is\s+', 'We identify a ', text)
    # Careful with "there is" at end of sentence — skip
    text = re.sub(r'\bIt is possible to\b', 'We can', text)
    
    # --- Redundant intensifiers ---
    text = re.sub(r'\bvery\s+', '', text)
    text = re.sub(r'\bextremely\s+', '', text)
    text = re.sub(r'\bremarkably\s+', '', text)
    text = re.sub(r'\bquite\s+', '', text)
    
    # --- "this work" vs "this paper" consistency ---
    # Sometimes "this work" is better than "this paper" for method descriptions
    # Already handled by earlier conversion, but ensure consistency
    
    # --- "aforementioned" / "above-mentioned" -> "above" or remove ---
    text = text.replace('aforementioned', 'above')
    
    # --- "In this section, we..." -> keep but trim surrounding fluff ---
    text = re.sub(r'\bIn what follows,?\s+', '', text)
    text = re.sub(r'\bIn the subsequent sections?\b', 'Below', text)
    
    # --- "We note that" -> keep only when genuinely noting ---
    # Leave this as-is; it's a valid academic construction
    
    # --- "broadly" / "generally" tightening ---
    # Keep them, they add appropriate hedging
    
    # --- Fix common phrasal redundancies ---
    text = re.sub(r'\b(?:both\s+)?together\s+with\b', 'with', text)
    text = re.sub(r'\beach and every\b', 'each', text)
    text = re.sub(r'\bfirst and foremost\b', 'first', text)
    
    # --- "the present study" -> "this study" ---
    text = re.sub(r'\bthe present study\b', 'this study', text)
    
    # --- "in terms of" -> rewrite ---
    text = re.sub(r'\bin terms of\b', 'in', text)
    
    return text


def remove_empty_headings(doc):
    """Remove empty heading paragraphs."""
    paras_to_remove = []
    for i, p in enumerate(doc.paragraphs):
        if 'Heading' in p.style.name and not p.text.strip():
            paras_to_remove.append(p)
    
    for p in paras_to_remove:
        p._element.getparent().remove(p._element)
    
    return len(paras_to_remove)


def compress_blank_lines(doc):
    """Remove excessive consecutive blank paragraphs."""
    blank_runs = 0
    to_remove = []
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            blank_runs += 1
            if blank_runs > 2:  # Max 2 consecutive blank lines
                to_remove.append(p)
        else:
            blank_runs = 0
    
    for p in to_remove:
        p._element.getparent().remove(p._element)
    
    return len(to_remove)


def clean_references_formatting(doc):
    """
    Apply consistent formatting to reference entries.
    Italicize book titles and journal names.
    """
    in_refs = False
    for i, p in enumerate(doc.paragraphs):
        if 'References' in p.text and 'Heading' in p.style.name:
            in_refs = True
            continue
        if in_refs and p.text.strip():
            text = p.text
            # Ensure hanging indent is properly applied
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(0)
            # Set font size
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'


def ensure_heading_formatting(doc):
    """Ensure consistent heading formatting."""
    for i, p in enumerate(doc.paragraphs):
        if 'Heading' in p.style.name:
            text = p.text.strip()
            if not text:
                continue
            for run in p.runs:
                # Set heading font
                pass  # Preserve existing heading styles


def fix_spacing_around_punctuation(s):
    """Fix spacing issues around punctuation, but preserve decimal numbers like 4.2.1."""
    # Remove space before comma/semicolon/colon
    s = re.sub(r'\s+,', ',', s)
    s = re.sub(r'\s+;', ';', s)
    s = re.sub(r'\s+:', ':', s)
    s = re.sub(r'\s+\.', '.', s)
    # Ensure space after comma/semicolon/colon if not end
    s = re.sub(r',(\S)', r', \1', s)
    s = re.sub(r';(\S)', r'; \1', s)
    s = re.sub(r':(\S)', r': \1', s)
    # Fix period spacing — but NOT for decimal/section numbers like "4.2.1" or "1.1"
    # Match period followed by non-space, non-digit (to preserve section numbering)
    s = re.sub(r'\.(\D)', r'. \1', s)
    # Remove double spaces
    s = re.sub(r'  +', ' ', s)
    return s


def fix_block_text(doc):
    """Update the block text (subtitle) formatting."""
    for i, p in enumerate(doc.paragraphs):
        if 'Block Text' in p.style.name:
            text = p.text.strip()
            if text:
                # Ensure consistent formatting
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                    run.font.italic = True
                # Center align
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def final_pass_refinements(text):
    """Final line-by-line refinements that are safe to apply broadly."""
    # Fix "cannot" -> "can not" or keep? Keep "cannot" — it's fine.
    
    # "i.e." and "e.g." spacing
    text = re.sub(r'\bi\.e\.', 'i.e.,', text)
    text = re.sub(r'\be\.g\.', 'e.g.,', text)
    
    # Fix "et al" -> "et al."
    text = re.sub(r'\bet al\b(?!\.)', 'et al.', text)
    
    # Fix "Fig." / "Table" consistency
    # Already fine
    
    # Ensure "Section X" cross-references have consistent formatting
    text = re.sub(r'\b[Ss]ection (\d+\.?\d*)', r'Section \1', text)
    
    return text


def main():
    doc = Document(SRC)
    
    # Step 1: Remove empty headings first
    removed_headings = remove_empty_headings(doc)
    print(f'Removed {removed_headings} empty heading paragraphs')
    
    # Step 2: Compress excessive blank lines
    removed_blanks = compress_blank_lines(doc)
    print(f'Removed {removed_blanks} excessive blank paragraphs')
    
    # Step 3: Apply text transformations paragraph by paragraph
    em_dash_count = 0
    polish_count = 0
    for i, p in enumerate(doc.paragraphs):
        original = p.text
        if not original.strip():
            continue
        
        text = original
        
        # 3a. Replace em dashes (different logic for headings vs body)
        is_heading = 'Heading' in p.style.name
        new_text = replace_em_dashes(text, is_heading=is_heading)
        if new_text != text:
            em_dash_count += text.count('—')
            text = new_text
        
        # 3b. Polish writing
        new_text = polish_writing(text)
        if new_text != text:
            polish_count += 1
            text = new_text
        
        # 3c. Final pass refinements
        new_text = final_pass_refinements(text)
        if new_text != text:
            text = new_text
        
        # 3d. Fix spacing around punctuation
        new_text = fix_spacing_around_punctuation(text)
        if new_text != text:
            text = new_text
        
        # Update paragraph if changed
        if text != original:
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = text
            else:
                p.add_run(text)
    
    print(f'Replaced {em_dash_count} em dashes')
    print(f'Polished {polish_count} paragraphs')
    
    # Step 4: Fix formatting
    clean_references_formatting(doc)
    fix_block_text(doc)
    
    doc.save(DST)
    print(f'Saved refined paper to: {DST}')


if __name__ == '__main__':
    main()
