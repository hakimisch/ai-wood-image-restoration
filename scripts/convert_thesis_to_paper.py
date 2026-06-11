#!/usr/bin/env python
"""
Convert thesis DOCX to journal-paper DOCX:
  - Remove "Chapter" designations
  - Replace "this thesis" with "this paper"/"this study"
  - Remove End-of-Chapter markers
  - Add References/Bibliography section
  - Condense verbose thesis-style passages
"""

import re
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r'C:\Hakimi\Internship\GUI Image Restoration\research\output\convertible.docx'
DST = r'C:\Hakimi\Internship\GUI Image Restoration\research\output\paper.docx'

def extract_all_paragraphs(doc):
    """Extract paragraphs with style info."""
    return [(i, p.style.name, p.text) for i, p in enumerate(doc.paragraphs)]

def fix_chapter_headings(text):
    """Remove 'Chapter' designation: 'Chapter 1: Introduction' -> '1. Introduction'
    Only transforms heading-level text (followed by colon), not inline references."""
    text = re.sub(r'\bChapter\s+(\d+):\s*', r'\1. ', text)
    # Handle "End of Chapter 1" -> (removed by fix_end_of_chapter)
    return text

def fix_inline_section_refs(text):
    """Convert inline 'Chapter X' references to 'Section X'."""
    text = re.sub(r'\bChapter\s+(\d+)\b', r'Section \1', text)
    return text

def fix_this_thesis(text):
    """Replace 'this thesis' with context-appropriate alternatives."""
    # "the thesis" -> "the paper" when referring to this work
    text = re.sub(r'\b[Tt]he thesis\b', r'this paper', text)
    # "This thesis" -> "This paper" (sentence start)
    text = re.sub(r'\bThis thesis\b', r'This paper', text)
    # "this thesis" -> "this paper" (lowercase)
    text = re.sub(r'\bthis thesis\b', r'this paper', text)
    # Handle "thesis" as standalone where it means "this work"
    text = re.sub(r'\bthesis\'s\b', r"paper's", text)
    return text

def fix_end_of_chapter(text):
    """Remove 'End of Chapter X' and 'End of Thesis Document' markers."""
    text = re.sub(r'End of Chapter\s+\d+.*$', '', text)
    text = re.sub(r'End of Thesis Document.*$', '', text)
    return text

def fix_thesis_document_tag(text):
    """Replace 'Thesis Document' -> 'Paper'."""
    text = text.replace('Thesis Document', 'Paper')
    return text

def fix_em_dashes_and_cleanup(text):
    """Clean up punctuation and references to thesis structure."""
    # Replace em dashes with spaced en dashes (more standard for papers)
    # Actually, let's keep em dashes since they're fine for papers
    # Just fix "this thesis" -> "this paper" patterns we might have missed
    text = text.replace('(this thesis)', '(this paper)')
    text = text.replace('(this thesis,', '(this paper,')
    return text

def condense_text(text):
    """Apply condensation to verbose thesis-style passages."""
    # Condensed version: remove some of the more verbose thesis-style framing
    # These are targeted replacements for very long-winded passages
    
    # Condense "Let us define..." -> "Define..."
    text = re.sub(r'\bLet us define\b', r'Define', text)
    
    # Condense "Let us consider" -> "Consider"
    text = re.sub(r'\bLet us consider\b', r'Consider', text)
    
    return text

def extract_citations(text):
    """Extract all unique in-text citations from text."""
    # Pattern: (Author, Year) or (Author et al., Year)
    citation_pattern = re.compile(r'\(([^()]*?\d{4}[a-z]?)\)')
    citations = set()
    
    # Also find "Author et al. (Year)" patterns
    named_pattern = re.compile(r'([A-Z][a-z]+(?:\s+et\s+al\.)?)\s*\((\d{4}[a-z]?)\)')
    
    for match in citation_pattern.findall(text):
        # Clean up: split multiple citations separated by ;
        parts = [p.strip() for p in match.split(';')]
        for part in parts:
            # Extract just author-year
            part = part.strip()
            if part and re.search(r'\d{4}', part):
                citations.add(part)
    
    for match in named_pattern.findall(text):
        author, year = match
        # Only add if it's an actual citation (not a number in parentheses)
        if len(year) >= 4:
            citations.add(f"{author} ({year})")
    
    return sorted(citations)

def build_references_section():
    """Build a References/Bibliography section from known citations in the text."""
    references = [
        ("Dong, C., Loy, C. C., He, K., & Tang, X. (2014). Learning a deep convolutional network for image super-resolution. In *European Conference on Computer Vision (ECCV)* (pp. 184–199). Springer.", "Dong et al., 2014"),
        ("Dosovitskiy, A., Beyer, L., Kolesnikov, A., et al. (2021). An image is worth 16x16 words: Transformers for image recognition at scale. In *International Conference on Learning Representations (ICLR)*.", "Dosovitskiy et al., 2021"),
        ("Goodfellow, I., Pouget-Abadie, J., Mirza, M., et al. (2014). Generative adversarial nets. In *Advances in Neural Information Processing Systems (NeurIPS)* (pp. 2672–2680).", "Goodfellow et al., 2014"),
        ("He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep residual learning for image recognition. In *IEEE Conference on Computer Vision and Pattern Recognition (CVPR)* (pp. 770–778).", "He et al., 2016"),
        ("Kim, J., Lee, J. K., & Lee, K. M. (2016). Accurate image super-resolution using very deep convolutional networks. In *IEEE Conference on Computer Vision and Pattern Recognition (CVPR)* (pp. 1646–1654).", "Kim et al., 2016"),
        ("Ledig, C., Theis, L., Huszár, F., et al. (2017). Photo-realistic single image super-resolution using a generative adversarial network. In *IEEE Conference on Computer Vision and Pattern Recognition (CVPR)* (pp. 4681–4690).", "Ledig et al., 2017"),
        ("Liang, J., Cao, J., Sun, G., et al. (2021). SwinIR: Image restoration using swin transformer. In *IEEE/CVF International Conference on Computer Vision Workshops (ICCVW)* (pp. 1833–1844).", "Liang et al., 2021"),
        ("Lucy, L. B. (1974). An iterative technique for the rectification of observed distributions. *The Astronomical Journal*, 79, 745.", "Lucy, 1974"),
        ("Richardson, W. H. (1972). Bayesian-based iterative method of image restoration. *Journal of the Optical Society of America*, 62(1), 55–59.", "Richardson, 1972"),
        ("Wang, Z., Bovik, A. C., Sheikh, H. R., & Simoncelli, E. P. (2004). Image quality assessment: From error visibility to structural similarity. *IEEE Transactions on Image Processing*, 13(4), 600–612.", "Wang et al., 2004"),
        ("Wang, X., Yu, K., Wu, S., et al. (2018). ESRGAN: Enhanced super-resolution generative adversarial networks. In *European Conference on Computer Vision (ECCV) Workshops* (pp. 63–79).", "Wang et al., 2018"),
        ("Wiener, N. (1949). *Extrapolation, Interpolation, and Smoothing of Stationary Time Series*. MIT Press.", "Wiener, 1949"),
    ]
    return references

def main():
    doc = Document(SRC)
    
    # Collect all paragraphs info
    paras = extract_all_paragraphs(doc)
    
    # Apply transformations paragraph by paragraph
    for i, p in enumerate(doc.paragraphs):
        original = p.text
        if not original.strip():
            continue
        
        text = original
        
        # Skip if this is a heading that starts a chapter
        # Actually let's process everything
        
        # Apply transformations
        # 1. Heading fix: "Chapter 1: Introduction" -> "1. Introduction"
        text = fix_chapter_headings(text)
        # 2. Remove "End of Chapter" markers BEFORE inline refs convert them
        text = fix_end_of_chapter(text)
        # 3. Inline refs: remaining "Chapter 4 presents" -> "Section 4 presents"
        text = fix_inline_section_refs(text)
        text = fix_this_thesis(text)
        text = fix_thesis_document_tag(text)
        text = condense_text(text)
        
        # Fix "remainder of this thesis" -> "remainder of this paper"
        text = text.replace('remainder of this thesis', 'remainder of this paper')
        
        # Fix "central thesis of this work" -> "central argument of this work"
        # (keep "thesis" when it means proposition/argument)
        
        # Clear the runs and set the new text
        if text != original:
            # Clear existing runs
            for run in p.runs:
                run.text = ''
            # Set text in first run
            if p.runs:
                p.runs[0].text = text
            else:
                p.add_run(text)
    
    # Now add References section at the end
    # Add a page break
    doc.add_paragraph()  # spacer
    
    # Add References heading
    heading = doc.add_heading('References', level=1)
    
    # Add references
    references = build_references_section()
    for ref_text, _ in references:
        para = doc.add_paragraph(ref_text)
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
        para.paragraph_format.space_after = Pt(6)
    
    doc.save(DST)
    print(f"Saved converted paper to: {DST}")

if __name__ == '__main__':
    main()
