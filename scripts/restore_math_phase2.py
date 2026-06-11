#!/usr/bin/env python
"""
Phase 2: Fix math variables that appear as single-space gaps (no double spaces).
The original LaTeX was completely stripped from the DOCX, leaving only single-space
gaps where variables like $I_{clear}$ or $\Phi_{phys}$ once were.
Strategy: pattern-match on neighboring English text.
"""

import re
from docx import Document

SRC = r'C:\Hakimi\Internship\GUI Image Restoration\research\output\paper_fixed.docx'
DST = r'C:\Hakimi\Internship\GUI Image Restoration\research\output\paper_fixed.docx'


def fix_math_vars(text):
    """Fix missing math variables using surrounding English text patterns."""
    
    # === Section 1.2 ===
    
    # "Let represent a latent" -> "Let I_clear represent a latent"
    text = re.sub(r'\bLet represent a latent\b',
                  'Let I_clear represent a latent', text)
    
    # "system produces an observation" (after "imaging" was mentioned)
    # But be careful not to match other "system produces" phrases
    text = re.sub(r'\bimaging system produces\b',
                  'imaging system Phi_phys produces', text)
    
    # "where is the microscope's true point spread function" -> "where h_PSF is..."
    text = re.sub(r"\bwhere is the microscope's true point spread function\b",
                  "where h_PSF is the microscope's true point spread function", text)
    
    # "interface, is read noise," -> "interface, eta_sensor is read noise,"
    text = re.sub(r'\binterface, is read noise,\b',
                  'interface, eta_sensor is read noise,', text)
    
    # "read noise, is the periodic intensity modulation"
    text = re.sub(r'\bread noise, is the periodic intensity modulation\b',
                  'read noise, eps_banding is the periodic intensity modulation', text)
    
    # "modulation, and is the radial falloff"
    text = re.sub(r'\bmodulation, and is the radial falloff\b',
                  'modulation, and eps_vignette is the radial falloff', text)
    
    # "where is a fixed isotropic Gaussian kernel" -> "where G_15x15 is a fixed..."
    text = re.sub(r'\bwhere is a fixed isotropic Gaussian kernel\b',
                  'where G_15x15 is a fixed isotropic Gaussian kernel', text)
    
    # "inverts , not " or "invert , not "
    text = re.sub(r'\binvert\s*, not\s*\b',
                  'invert Phi_simple, not Phi_phys', text)
    text = re.sub(r'\binverts\s*, not\s*\b',
                  'inverts Phi_simple, not Phi_phys', text)
    
    # === Section 1.4 (VoL) ===
    
    # "where is the discrete Laplacian approximation using a kernel"
    text = re.sub(r'\bwhere is the discrete Laplacian approximation using a kernel\b',
                  'where L is the discrete Laplacian approximation using a 3x3 kernel', text)
    
    # === Section 2.2 (receptive field) ===
    
    # "of a -layer CNN with kernels and stride grows linearly as"
    text = re.sub(r'\bof a -layer CNN with kernels and stride grows linearly as\b',
                  'of a D-layer CNN with K x K kernels and stride S=1 grows linearly as R = 1 + D(K-1)', text)
    
    # === Section 3.3 (model architecture) ===
    
    # "of only pixels" -> "of only 7x7 pixels"
    text = re.sub(r'\bof only pixels\b',
                  'of only 7x7 pixels', text)
    
    # "achieving a receptive field" -> "achieving a 17x17 receptive field" (near kernel mentions)
    # More specific: after discussing 9-5-5 kernel cascade
    text = re.sub(r'\b9-5-5 kernel cascade achieving a receptive field\b',
                  '9-5-5 kernel cascade achieving a 17x17 receptive field', text)
    
    # "a convolution mapping 3 RGB channels" -> "a 3x3 convolution..."
    text = re.sub(r'\ba convolution mapping 3 RGB channels\b',
                  'a 3x3 convolution mapping 3 RGB channels', text)
    
    # === Section 3.4.1 (loss function) ===
    
    # "where is the mean squared error" -> "where L_MSE is the mean squared error"
    text = re.sub(r'\bwhere is the mean squared error\b',
                  'where L_MSE is the mean squared error', text)
    
    # "and is the mean absolute error" -> "and L_L1 is the mean absolute error"
    text = re.sub(r'\band is the mean absolute error\b',
                  'and L_L1 is the mean absolute error', text)
    
    # === Section 4.3.1 (MSE norms) ===
    
    # "minimise the expected norm" -> "minimise the expected L2 norm"
    text = re.sub(r'\bminimise the expected norm\b',
                  'minimise the expected L2 norm', text)
    
    # === Section 4.3.4 (TV loss) ===
    
    # "penalises the norm of the image gradient" -> "penalises the L1 norm..."
    text = re.sub(r'\bpenalises the norm of the image gradient\b',
                  'penalises the L1 norm of the image gradient', text)
    
    # "weights ranging from to" -> "weights ranging from 10^-4 to 10^-2"
    text = re.sub(r'\bweights ranging from to\b',
                  'weights ranging from 10^-4 to 10^-2', text)
    
    return text


def fix_3_6_semicolons(text):
    """Fix remaining semicolons in Section 3.6 compact items."""
    if text.startswith('Camera ISP'):
        text = re.sub(r'\(gamma correction\);\s*[Ee]arly',
                      '(gamma correction): early', text)
    if text.startswith('LED banding'):
        text = re.sub(r'[Bb]anding;\s*[Mm]odels',
                      'banding: models', text)
    if text.startswith('Sensor noise'):
        text = re.sub(r'[Nn]oise;\s*[Rr]ead',
                      'noise: read', text)
    return text


def main():
    doc = Document(SRC)
    fixes = 0
    
    for i, p in enumerate(doc.paragraphs):
        original = p.text
        if not original.strip():
            continue
        
        text = original
        
        # Apply math fixes
        new_text = fix_math_vars(text)
        if new_text != text:
            text = new_text
        
        # Apply semicolon fixes (Section 3.6)
        new_text = fix_3_6_semicolons(text)
        if new_text != text:
            text = new_text
        
        if text != original:
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = text
            else:
                p.add_run(text)
            fixes += 1
    
    doc.save(DST)
    print(f'Fixed {fixes} paragraphs')
    
    # Verify
    print('\n=== Verification ===')
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        
        # Check key fixes
        if 'Let I_clear' in t:
            print(f'  [OK] Section 1.2: Let I_clear')
        if 'Phi_phys' in t:
            print(f'  [OK] Section 1.2: Phi_phys')
        if 'h_PSF' in t:
            print(f'  [OK] Section 1.2: h_PSF')
        if 'eta_sensor' in t:
            print(f'  [OK] Section 1.2: eta_sensor')
        if 'eps_banding' in t:
            print(f'  [OK] Section 1.2: eps_banding')
        if 'eps_vignette' in t:
            print(f'  [OK] Section 1.2: eps_vignette')
        if 'G_15x15' in t:
            print(f'  [OK] Section 1.2: G_15x15')
        if 'Phi_simple' in t:
            print(f'  [OK] Section 1.2: Phi_simple')
        if 'L is the discrete Laplacian' in t:
            print(f'  [OK] Section 1.4: VoL')
        if 'D-layer CNN' in t:
            print(f'  [OK] Section 2.2: receptive field')
        if 'L_MSE is the mean squared error' in t:
            print(f'  [OK] Section 3.4.1: loss function')
        if 'L2 norm' in t:
            print(f'  [OK] Section 4.3.1: L2 norm')
        if 'L1 norm of the image gradient' in t:
            print(f'  [OK] Section 4.3.4: L1 norm')
        if '10^-4 to 10^-2' in t:
            print(f'  [OK] Section 4.3.4: TV weights')
        if 'Camera ISP (gamma correction):' in t:
            print(f'  [OK] Section 3.6: Camera ISP')


if __name__ == '__main__':
    main()
