#!/usr/bin/env python
"""
Repair missing math variables, broken equation descriptions, and Section 3.6
punctuation in paper.docx.

All LaTeX math variables sourced from research_paper_draft.md.
Uses placeholders (e.g., "L_2 norm") since DOCX cannot render LaTeX math natively.
"""

import re
from docx import Document

SRC = r'C:\Hakimi\Internship\GUI Image Restoration\research\output\paper.docx'
DST = r'C:\Hakimi\Internship\GUI Image Restoration\research\output\paper_fixed.docx'


def _sub(regex, replacement, text):
    """Safe re.sub using lambda to avoid backslash escaping issues."""
    return re.sub(regex, lambda m: replacement, text)


def fix_section_1_2(text):
    """Restore stripped math variables in Section 1.2."""
    # "Let  represent a latent" -> "Let I_clear represent a latent"
    text = re.sub(r'Let\s{2,}represent a latent',
                  lambda m: 'Let I_clear represent a latent', text)
    
    # "physical imaging system  produces" -> "physical imaging system Phi produces"
    text = re.sub(r'physical imaging system\s{2,}produces',
                  lambda m: 'physical imaging system Phi_phys produces', text)
    
    # "where  is the microscope's true point spread function"
    text = re.sub(r"where\s{2,}is the microscope's true point spread function",
                  lambda m: "where h_PSF is the microscope's true point spread function", text)
    
    # "interface,  is read noise"
    text = re.sub(r'interface,\s{2,}is read noise',
                  lambda m: 'interface, eta_sensor is read noise', text)
    
    # "read noise,  is the periodic intensity modulation"
    text = re.sub(r'read noise,\s{2,}is the periodic intensity modulation',
                  lambda m: 'read noise, eps_banding is the periodic intensity modulation', text)
    
    # "modulation,  and  is the radial falloff"
    text = re.sub(r'modulation,\s{2,}and\s{2,}is the radial falloff',
                  lambda m: 'modulation, and eps_vignette is the radial falloff', text)
    
    # "where  is a fixed isotropic Gaussian kernel"
    text = re.sub(r'where\s{2,}is a fixed isotropic Gaussian kernel',
                  lambda m: 'where G_15x15 is a fixed isotropic Gaussian kernel', text)
    
    # "invert  , not "
    text = re.sub(r'invert\s{2,}, not\s{2,}',
                  lambda m: 'invert Phi_simple, not Phi_phys', text)
    
    return text


def fix_section_1_4(text):
    """Restore math in Section 1.4 (VoL equation description)."""
    # "where  is the discrete Laplacian"
    text = re.sub(r'where\s{2,}is the discrete Laplacian',
                  lambda m: 'where L is the discrete Laplacian', text)
    
    # "using a  kernel"
    text = re.sub(r'using a\s{2,}kernel',
                  lambda m: 'using a 3x3 kernel', text)
    
    return text


def fix_section_2_2(text):
    """Restore receptive field math in Section 2.2."""
    text = re.sub(r'receptive field of a\s{2,}layer CNN with\s{2,}kernels',
                  lambda m: 'receptive field of a D-layer CNN with K x K kernels', text)
    
    text = re.sub(r'and stride\s{2,}grows linearly as\s{2,}',
                  lambda m: 'and stride S=1 grows linearly as R = 1 + D(K-1)', text)
    
    return text


def fix_section_3_6(text):
    """Fix Section 3.6: replace semicolons with colons in compact list items."""
    if text.startswith('Camera ISP'):
        text = re.sub(r'\(gamma correction\);\s*Early', '(gamma correction): early', text)
    if text.startswith('LED banding'):
        text = re.sub(r'Banding;\s*Models', 'Banding: models', text)
    if text.startswith('Sensor noise'):
        text = re.sub(r'noise;\s*Read noise', 'noise: read noise', text)
    return text


def fix_section_3_4(text):
    """Restore math in Section 3.4.1 (loss function)."""
    text = re.sub(r'where\s{2,}is the mean squared error',
                  lambda m: 'where L_MSE is the mean squared error (nn.MSELoss())', text)
    
    text = re.sub(r'and\s{2,}is the mean absolute error',
                  lambda m: 'and L_L1 is the mean absolute error (nn.L1Loss())', text)
    
    return text


def fix_section_4(text):
    """Restore math in Section 4.x."""
    text = re.sub(r'minimise the expected\s{2,}norm',
                  lambda m: 'minimise the expected L2 norm', text)
    
    text = re.sub(r'penalises the\s{2,}norm of the image gradient',
                  lambda m: 'penalises the L1 norm of the image gradient', text)
    
    text = re.sub(r'weights ranging from\s{2,}to',
                  lambda m: 'weights ranging from 10^-4 to 10^-2', text)
    
    return text


def fix_section_3_3(text):
    """Restore math in Section 3.3 (model architecture)."""
    text = re.sub(r'receptive field of only\s{2,}pixels',
                  lambda m: 'receptive field of only 7x7 pixels', text)
    
    text = re.sub(r'achieving a\s{2,}receptive field',
                  lambda m: 'achieving a 17x17 receptive field', text)
    
    text = re.sub(r'receptive field of\s{2,}pixels\)',
                  lambda m: 'receptive field of 41x41 pixels)', text)
    
    text = re.sub(r'a\s{2,}convolution mapping 3 RGB channels',
                  lambda m: 'a 3x3 convolution mapping 3 RGB channels', text)
    
    return text


def main():
    doc = Document(SRC)
    fixes = 0
    
    for i, p in enumerate(doc.paragraphs):
        original = p.text
        if not original.strip():
            continue
        
        text = original
        for fn in [fix_section_1_2, fix_section_1_4, fix_section_2_2,
                   fix_section_3_3, fix_section_3_4, fix_section_3_6,
                   fix_section_4]:
            text = fn(text)
        
        if text != original:
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = text
            else:
                p.add_run(text)
            fixes += 1
    
    doc.save(DST)
    print(f'Restored math in {fixes} paragraphs')
    
    # Verify
    print('\n=== Spot checks ===')
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        if 'Let I_clear' in t:
            print(f'  Section 1.2 OK: ...{t[:120]}')
        if 'Phi_phys produces' in t:
            print(f'  Section 1.2 OK: ...{t[re.search("Phi_phys", t).start()-20:re.search("Phi_phys", t).start()+20]}...')
        if 'h_PSF' in t and 'microscope' in t:
            print(f'  Section 1.2 PSF OK')
        if 'D-layer CNN' in t:
            print(f'  Section 2.2 OK: {t[t.find("receptive field"):t.find("receptive field")+100]}')
        if 'Camera ISP (gamma correction):' in t:
            print(f'  Section 3.6 OK: {t[:100]}')
        if 'L2 norm' in t:
            print(f'  Section 4.3.1 OK')
        if 'L1 norm of the image gradient' in t:
            print(f'  Section 4.3.4 OK')
        if 'L_MSE is the mean squared error' in t:
            print(f'  Section 3.4.1 OK')
    
    # Check for remaining gaps
    print('\n=== Remaining gaps check ===')
    gaps = 0
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        m = re.search(r'(?<=[a-z])\s{3,}(?=[a-z])', t)
        if m:
            gaps += 1
            start = m.start()
            ctx = t[max(0, start-15):start+40]
            print(f'  Gap [{i}]: ...{ctx}...')
    print(f'  Total remaining gaps: {gaps}')


if __name__ == '__main__':
    main()
