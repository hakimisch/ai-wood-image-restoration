#!/usr/bin/env python
"""
Repair the paper DOCX after botched spacing fix:
  - Fix "1. 1" -> "1.1" in section numbers
  - Fix "Phase 1; The Gaussian-Only Catastrophe" -> "Phase 1: The Gaussian-Only Catastrophe" (subtitle semicolons)
  - Fix "0. 8" -> "0.8" in decimal numbers
  - Fix "4. 2. 1" -> "4.2.1" in deeply nested headings
  - Fix other spacing artifacts
"""

import re
import docx
from docx import Document

SRC = r'C:\Hakimi\Internship\GUI Image Restoration\research\output\paper.docx'
DST = r'C:\Hakimi\Internship\GUI Image Restoration\research\output\paper.docx'

def fix_section_numbers(text):
    """
    Fix section numbering: "1. 1" -> "1.1", "3. 2. 2" -> "3.2.2"
    Matches patterns like digit(s). space digit(s)
    """
    # Fix "digit. digit" inside section numbers (e.g., "3. 2. 2" -> "3.2.2")
    # This pattern: digit(s) followed by period-space-digit(s)
    text = re.sub(r'(\d+)\. (\d)', r'\1.\2', text)
    return text

def fix_subtitle_semicolons(text):
    """
    Fix heading subtitles where em dash was replaced with semicolon:
    "Phase 1; The Gaussian-Only Catastrophe" -> "Phase 1: The Gaussian-Only Catastrophe"
    Only applies to heading-like patterns (short left side, capitalized right side).
    """
    # Pattern: phrase ending with digit or short word, semicolon, space, Capitalized word
    # Only fix if the left side looks like a subtitle label
    text = re.sub(r'(\b(?:Phase|Stage|Part|Section|Module)\s+\d+);\s+([A-Z])', r'\1: \2', text)
    # Also fix "4.3.4 Total Variation Loss; A Cautionary Tale" -> colon
    text = re.sub(r'(\.\d+ [A-Za-z ]+); ([A-Z])', r'\1: \2', text)
    return text

def fix_decimal_spaces(text):
    """
    Fix decimal numbers: "0. 8" -> "0.8", "1. 5" -> "1.5"
    """
    text = re.sub(r'(\d+)\. (\d+)', r'\1.\2', text)
    return text

def main():
    doc = Document(SRC)
    
    fix_count = 0
    for i, p in enumerate(doc.paragraphs):
        original = p.text
        if not original.strip():
            continue
        
        text = original
        
        # Apply fixes in order
        text = fix_section_numbers(text)
        text = fix_subtitle_semicolons(text)
        text = fix_decimal_spaces(text)
        
        # Also fix any remaining "3. 2. 2" triple patterns
        text = re.sub(r'(\d+)\. (\d+)\. (\d+)', r'\1.\2.\3', text)
        
        # Fix double spaces
        text = re.sub(r'  +', ' ', text)
        
        if text != original:
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = text
            else:
                p.add_run(text)
            fix_count += 1
    
    doc.save(DST)
    print(f'Fixed {fix_count} paragraphs')

if __name__ == '__main__':
    main()
